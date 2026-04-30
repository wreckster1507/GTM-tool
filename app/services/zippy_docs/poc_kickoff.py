"""PoC Kickoff generator — fills Beacon's Drive PoC Kickoff template.

Same flow as the MOM generator: find the indexed Drive template, export
the Google Doc to .docx, hand the structure to Claude with the
``poc_kickoff`` doc-type instruction set, patch rewrites back, upload as
a new editable Google Doc.

The template is a Google Doc with bracketed placeholders ([Client Name],
[Insert URL], [Insert Password], etc.) plus prose blocks for use cases
and deliverables. Zippy's job is to replace bracketed slots with values
from the user's Gmail thread context, leave headings untouched, and keep
the template's two use-case structure even when only one was discussed.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from sqlmodel import select as sm_select

from app.clients.anthropic_client import get_anthropic_client
from app.clients.google_drive import upload_as_google_doc
from app.config import settings
from app.database import AsyncSessionLocal as async_session
from app.models.user_email_connection import UserEmailConnection
from app.models.zippy import IndexedDriveFile
from app.services.zippy_docs.base import (
    GeneratedDocument,
    build_output_path,
    human_today,
)
from app.services.zippy_docs.claude_rewriter import rewrite_with_claude
from app.services.zippy_docs.doc_rewriter import (
    extract_docx_structure,
    rewrite_docx_content,
)

logger = logging.getLogger(__name__)


@dataclass
class PoCKickoffInput:
    """Inputs for the PoC Kickoff generator.

    ``email_thread_content`` is the source of truth for everything the
    rewriter needs (use cases, credentials, dates). It's clipped at 20k
    chars before being handed to Claude so a runaway forward-chain
    doesn't blow the context window.
    """
    client_name: str
    email_thread_content: str = ""
    meeting_date: Optional[str] = None
    prepared_by: Optional[str] = None
    extra_context: Optional[str] = None


# ── Drive template discovery ────────────────────────────────────────────────

_POC_KICKOFF_KEYWORDS = [
    "template of beacon poc kickoff",
    "beacon poc kickoff template",
    "poc kickoff template",
    "beacon poc kickoff",
    "poc kickoff",
]

# Per-file in-memory cache; keyed by drive_modified_at so a re-uploaded
# template invalidates automatically. Same shape as the MOM cache.
_TEMPLATE_CACHE: dict[str, tuple[str, bytes]] = {}


async def _find_poc_kickoff_template_row(
    user_id: Optional[str] = None,
) -> Optional[IndexedDriveFile]:
    """Return the IndexedDriveFile row for the PoC Kickoff template, or None.

    Scope rules: prefer the user's own synced file, fall back to admin/shared,
    never leak another user's private file.
    """
    from sqlalchemy import or_

    async with async_session() as session:
        for keyword in _POC_KICKOFF_KEYWORDS:
            stmt = sm_select(IndexedDriveFile).where(
                IndexedDriveFile.name.ilike(f"%{keyword}%"),
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        IndexedDriveFile.owner_user_id == user_id,
                        IndexedDriveFile.is_admin == True,  # noqa: E712
                    )
                )
            else:
                stmt = stmt.where(IndexedDriveFile.is_admin == True)  # noqa: E712

            if user_id is not None:
                from sqlalchemy import case
                user_priority = case(
                    (IndexedDriveFile.owner_user_id == user_id, 0),
                    else_=1,
                )
                stmt = stmt.order_by(
                    user_priority,
                    IndexedDriveFile.last_indexed_at.desc(),
                )
            else:
                stmt = stmt.order_by(IndexedDriveFile.last_indexed_at.desc())

            row = (await session.execute(stmt.limit(1))).scalar_one_or_none()
            if row:
                logger.info(
                    "PoC Kickoff template found via keyword=%r: name=%s "
                    "file_id=%s mime=%s owner=%s is_admin=%s",
                    keyword, row.name, row.drive_file_id, row.mime_type,
                    row.owner_user_id, row.is_admin,
                )
                return row
    logger.info("No PoC Kickoff template matched for user_id=%s", user_id)
    return None


async def _fetch_template_bytes(
    user_id: Optional[str] = None,
) -> tuple[Optional[bytes], Optional[str]]:
    """Fetch the PoC Kickoff template bytes for the given user.

    The template is published as a Google Doc, so the typical path is
    Drive's ``/export?mimeType=...docx`` endpoint. Falls back to a direct
    download if a future re-upload makes it a real .docx.
    """
    row = await _find_poc_kickoff_template_row(user_id=user_id)
    if not row:
        return None, (
            "No PoC Kickoff template indexed. Upload a file named "
            "'Template of Beacon PoC Kickoff' to your synced Drive folder "
            "and click Sync."
        )

    file_id = row.drive_file_id
    actual_mime = row.mime_type
    cache_key = (
        row.drive_modified_at.isoformat() if row.drive_modified_at else "no-modified"
    )

    cached = _TEMPLATE_CACHE.get(file_id)
    if cached and cached[0] == cache_key:
        logger.info(
            "PoC Kickoff template cache HIT for %s (%d bytes)",
            file_id, len(cached[1]),
        )
        return cached[1], None

    from app.clients import google_drive

    async with async_session() as session:
        result = await session.execute(
            sm_select(UserEmailConnection).where(
                UserEmailConnection.user_id == row.owner_user_id,
                UserEmailConnection.is_active == True,  # noqa: E712
            )
        )
        connection = result.scalar_one_or_none()

        if not connection:
            logger.warning(
                "No active Drive connection for owner_user_id=%s — "
                "cannot fetch PoC Kickoff template",
                row.owner_user_id,
            )
            return None, (
                f"Drive connection for the owner of '{row.name}' is "
                "inactive. Reconnect Google Drive in Settings and re-sync."
            )

        try:
            if actual_mime == "application/vnd.google-apps.document":
                # Mirror the MOM Google-Doc export branch verbatim — same
                # endpoint, same headers, same target mime so the LLM gets
                # a real .docx structure to walk.
                import httpx
                from app.clients.google_drive import (
                    _ensure_token,
                    DRIVE_API_BASE,
                )
                access_token, _updated_token = await _ensure_token(
                    connection.token_data,
                    settings.gmail_client_id,
                    settings.gmail_client_secret,
                )
                docx_mime = (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
                async with httpx.AsyncClient(timeout=60) as http:
                    resp = await http.get(
                        f"{DRIVE_API_BASE}/files/{file_id}/export",
                        headers={"Authorization": f"Bearer {access_token}"},
                        params={"mimeType": docx_mime},
                    )
                if resp.status_code != 200:
                    return None, (
                        f"Drive export to .docx failed for '{row.name}' "
                        f"(status {resp.status_code}): {resp.text[:300]}"
                    )
                data = resp.content
            else:
                data, _mime, _updated = await google_drive.download_file_bytes(
                    file_id=file_id,
                    mime_type=actual_mime,
                    token_data=connection.token_data,
                    client_id=settings.gmail_client_id,
                    client_secret=settings.gmail_client_secret,
                )
            if not data:
                return None, (
                    f"Drive returned empty bytes for {row.name} "
                    f"(file_id={file_id}, mime={actual_mime})."
                )
            if not data.startswith(b"PK"):
                return None, (
                    f"'{row.name}' didn't come back as a valid .docx "
                    f"(first bytes: {data[:8]!r})."
                )
            logger.info(
                "PoC Kickoff template downloaded fresh: %d bytes", len(data)
            )
            _TEMPLATE_CACHE[file_id] = (cache_key, data)
            return data, None
        except Exception as exc:
            logger.exception(
                "Failed to download PoC Kickoff template file_id=%s: %s",
                file_id, exc,
            )
            return None, (
                f"Drive download failed for {row.name}: "
                f"{type(exc).__name__}: {exc}"
            )


# ── Inspection ──────────────────────────────────────────────────────────────


async def inspect_poc_kickoff_template(
    user_id: Optional[str] = None,
) -> dict:
    """Report whether the PoC Kickoff template is reachable + section count.

    Same return shape as ``inspect_mom_template`` so the agent's tool
    executor can format the response identically.
    """
    template_bytes, err = await _fetch_template_bytes(user_id=user_id)
    if not template_bytes:
        return {
            "found": False,
            "error": err or "Unknown error fetching PoC Kickoff template.",
            "section_count": 0,
            "sections": [],
        }

    structure = extract_docx_structure(template_bytes)
    content_blocks = [
        b for b in structure
        if not b["is_structural"] and b.get("text", "").strip()
    ]
    preview = [
        {"index": b["block_index"], "text": (b.get("text") or "")[:120]}
        for b in content_blocks[:10]
    ]

    row = await _find_poc_kickoff_template_row(user_id=user_id)
    return {
        "found": True,
        "template_name": row.name if row else None,
        "template_drive_file_id": row.drive_file_id if row else None,
        "section_count": len(content_blocks),
        "sections": preview,
    }


# ── Fallback renderer (no template available) ──────────────────────────────


def _render_fallback_docx(data: PoCKickoffInput, path) -> None:
    """Build a basic PoC Kickoff .docx when the Drive template is missing.

    Mirrors the canonical template's section order so the AE still gets
    something usable, but with a visible banner so it's clear this is a
    fallback rather than the brand-consistent Drive version.
    """
    doc = Document()
    title = doc.add_heading(
        f"PoC Kickoff — {data.client_name}", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)

    banner = doc.add_paragraph()
    br = banner.add_run(
        "DRAFT — Beacon's PoC Kickoff template was not reachable in Drive. "
        "Add it to your indexed folder and re-sync to use the official "
        "version."
    )
    br.italic = True
    br.font.size = Pt(9)
    br.font.color.rgb = RGBColor(0xB2, 0x00, 0x00)

    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run(data.meeting_date or human_today())
    meta.add_run("\n")
    meta.add_run("Prepared by: ").bold = True
    meta.add_run(data.prepared_by or "Beacon")

    def heading(text: str) -> None:
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)

    heading("Objective")
    doc.add_paragraph(
        f"Run a Proof of Concept for {data.client_name} to validate "
        "Beacon's automation against the workflow discussed."
    )

    heading("Login Credentials")
    doc.add_paragraph("URL: [Insert URL]")
    doc.add_paragraph("Username: [Insert Email/ID]")
    doc.add_paragraph("Password: [Insert Password]")

    heading("Use Cases")
    doc.add_paragraph("Use Case 1: [Use Case 1 title]")
    doc.add_paragraph("Business Problem: [Business Problem 1]")
    doc.add_paragraph("Expected Outcome: [Expected Outcome 1]")
    doc.add_paragraph()
    doc.add_paragraph("Use Case 2: [Use Case 2 title]")
    doc.add_paragraph("Business Problem: [Business Problem 2]")
    doc.add_paragraph("Expected Outcome: [Expected Outcome 2]")

    heading("Deliverables")
    doc.add_paragraph("Deliverable 1: [Deliverable 1 focus]")
    doc.add_paragraph("Deliverable 2: [Deliverable 2 focus]")

    heading("Timeline")
    doc.add_paragraph("Kickoff: [Insert Start Date]")
    doc.add_paragraph("Expected completion: [Insert End Date]")

    heading("Next Steps")
    if data.extra_context:
        doc.add_paragraph(data.extra_context)
    else:
        doc.add_paragraph(
            "Confirm credentials, schedule kickoff, and align on success "
            "criteria before the first working session."
        )

    footer = doc.add_paragraph()
    fr = footer.add_run(
        f"\nGenerated by Zippy on "
        f"{datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    doc.save(str(path))


# ── Public entry points ─────────────────────────────────────────────────────


async def generate(
    data: PoCKickoffInput,
    user_id: Optional[str] = None,
) -> GeneratedDocument:
    """Generate a PoC Kickoff doc by rewriting the Drive template with email context."""
    path, url = build_output_path("PoC_Kickoff", data.client_name)
    template_bytes, err = await _fetch_template_bytes(user_id=user_id)
    source = "template"
    rewritten: list[dict] = []  # populated only on the template path; the
    #                              fallback renderer doesn't go through Claude.

    if not template_bytes:
        logger.warning(
            "PoC Kickoff template unavailable (%s) — rendering fallback docx",
            err,
        )
        _render_fallback_docx(data, path)
        source = "fallback"
    else:
        structure = extract_docx_structure(template_bytes)

        user_inputs = {
            "client_name": data.client_name,
            "meeting_date": data.meeting_date or human_today(),
            # Don't paper over a missing AE name with a brand string —
            # leaving the placeholder makes it visibly unfilled, which is
            # the right cue for the AE to fill it in before sending.
            "prepared_by": data.prepared_by or "[Your Name]",
            # Cap email content so a long forward-chain doesn't blow the
            # context window. 20k chars is ~5k tokens — plenty for a PoC
            # discussion thread without crowding out the structure.
            "email_content": (data.email_thread_content or "")[:20000],
            "extra_context": data.extra_context or "",
        }

        client = get_anthropic_client()
        rewritten = await rewrite_with_claude(
            structure=structure,
            user_inputs=user_inputs,
            doc_type="poc_kickoff",
            client=client,
            model=settings.CLAUDE_MODEL_STANDARD,
        )

        out_bytes = rewrite_docx_content(template_bytes, rewritten)
        path.write_bytes(out_bytes)

    # Plain-text dump of the rewritten body so generate_poc_ppt can read
    # the kickoff content off the tool result without us having to add a
    # separate "fetch a Drive doc" tool. Empty rewrites are skipped so we
    # don't blow up on the NEXT STEPS section (deliberately blanked).
    body_text = "\n".join(
        (item.get("new_text") or "").strip()
        for item in (rewritten or [])
        if (item.get("new_text") or "").strip()
    )

    result = GeneratedDocument(
        filename=path.name,
        path=str(path),
        url=url,
        kind="poc_kickoff",
        summary=(
            f"PoC Kickoff drafted from {source} for {data.client_name}."
        ),
        created_at=datetime.utcnow(),
        body_text=body_text,
    )

    await _try_upload_to_drive(
        result, path, data.client_name, user_id=user_id
    )
    return result


async def _try_upload_to_drive(
    doc: GeneratedDocument,
    path,
    client_name: str,
    *,
    user_id: Optional[str],
) -> None:
    """Upload the filled .docx to Drive as an editable Google Doc.

    Best-effort, silently skipped on permission/upload failure — same
    contract as ``mom._try_upload_to_drive``.
    """
    if doc.drive_url:
        return
    if user_id:
        from app.services.zippy_docs.base import (
            cache_upload,
            get_cached_upload,
        )
        cached = get_cached_upload(str(user_id), client_name, doc.kind)
        if cached:
            doc.drive_url = cached
            logger.info(
                "Reusing cached PoC Kickoff Drive upload: %s", cached
            )
            return
    try:
        async with async_session() as session:
            from sqlalchemy import or_, case as sa_case

            stmt = sm_select(UserEmailConnection).where(
                UserEmailConnection.is_active == True,  # noqa: E712
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        UserEmailConnection.user_id == user_id,
                        UserEmailConnection.is_admin_folder == True,  # noqa: E712
                    )
                )
                priority = sa_case(
                    (UserEmailConnection.user_id == user_id, 0), else_=1
                )
                stmt = stmt.order_by(priority)
            result = await session.execute(stmt.limit(1))
            connection = result.scalar_one_or_none()

        if not connection:
            logger.info(
                "No active Drive connection — skipping PoC Kickoff upload"
            )
            return

        docx_bytes = path.read_bytes()
        gdoc_name = (
            f"PoC Kickoff — {client_name} — "
            f"{doc.created_at.strftime('%d %b %Y')}"
        )
        folder_id = connection.selected_drive_folder_id or None

        file_id, web_view_link = await upload_as_google_doc(
            filename=gdoc_name,
            docx_bytes=docx_bytes,
            token_data=connection.token_data,
            client_id=settings.gmail_client_id,
            client_secret=settings.gmail_client_secret,
            parent_folder_id=folder_id,
        )
        doc.drive_file_id = file_id
        doc.drive_url = web_view_link
        logger.info(
            "PoC Kickoff uploaded to Google Docs: %s", web_view_link
        )
        if user_id and web_view_link:
            from app.services.zippy_docs.base import cache_upload
            cache_upload(
                str(user_id), client_name, doc.kind, web_view_link
            )
    except PermissionError as exc:
        logger.info(
            "drive.file scope not yet granted — skipping PoC Kickoff "
            "upload: %s", exc,
        )
    except Exception as exc:
        logger.warning(
            "Google Docs upload failed for PoC Kickoff (non-fatal): %s",
            exc,
        )
