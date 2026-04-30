"""NDA generator — rewrites Beacon's Drive NDA template using user inputs.

The previous version scanned the template for ``______`` blanks and asked the
agent to fill them by index. That broke the moment a template was re-styled
or a blank was merged/split. The new flow is template-structure agnostic:

  1. ``inspect_template(jurisdiction)`` downloads the jurisdiction's template
     and reports how many content sections it contains, so the agent can
     confirm the template is reachable before collecting party details.

  2. ``generate(data)`` downloads the template, extracts every paragraph as a
     structured block, asks Claude to rewrite only the non-structural content
     from the user's inputs (parties, dates, jurisdiction, purpose, term),
     and patches the rewrites back preserving formatting. If Drive is
     unavailable, we fall back to a synthetic draft with a visible banner.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Literal, Optional

from docx import Document
from sqlmodel import select as sm_select

from app.clients.anthropic_client import get_anthropic_client
from app.clients.google_drive import upload_as_google_doc
from app.config import settings
from app.database import AsyncSessionLocal as async_session
from app.models.user_email_connection import UserEmailConnection
from app.models.zippy import IndexedDriveFile
from app.services.zippy_docs.base import (
    GeneratedDocument,
    build_output_path,
)
from app.services.zippy_docs.claude_rewriter import rewrite_with_claude
from app.services.zippy_docs.doc_rewriter import (
    extract_docx_structure,
    rewrite_docx_content,
)

logger = logging.getLogger(__name__)

Jurisdiction = Literal["india", "us", "singapore"]


@dataclass
class NDAInput:
    """Inputs for the NDA generator.

    ``fills`` is kept for backwards compatibility with older callers, but it
    is no longer interpreted positionally. It's now just a free-form dict
    passed through to Claude alongside the named fields — useful when the
    agent wants to forward extra details (e.g. registered office address).
    """
    jurisdiction: Jurisdiction
    receiving_party: Optional[str] = None
    disclosing_party: Optional[str] = None
    mutual: Optional[bool] = None
    purpose: Optional[str] = None
    term_years: Optional[int] = None
    effective_date: Optional[str] = None
    governing_city: Optional[str] = None
    extra_clauses: list[str] = field(default_factory=list)
    fills: dict[str, str] = field(default_factory=dict)


_JURISDICTION_LABEL = {
    "india": "Republic of India",
    "us": "State of Delaware, United States of America",
    "singapore": "Republic of Singapore",
}


def _template_drive_id(jurisdiction: Jurisdiction) -> str:
    return {
        "india": settings.NDA_TEMPLATE_DRIVE_ID_INDIA,
        "us": settings.NDA_TEMPLATE_DRIVE_ID_US,
        "singapore": settings.NDA_TEMPLATE_DRIVE_ID_SINGAPORE,
    }[jurisdiction]


# ── Drive template fetch ─────────────────────────────────────────────────────
#
# Auto-discovery first (mirrors mom._find_mom_template_row): scan
# IndexedDriveFile for filename matches, preferring the user's own file then
# admin-shared. Only falls back to env-var file_id if nothing is indexed.

# General keywords tried for every jurisdiction. Order matters — most specific
# first so "Template of Beacon NDA" beats a generic "NDA" hit.
_NDA_GENERAL_KEYWORDS = [
    "template of beacon nda",
    "beacon nda",
    "nda template",
    "nda",
]

_NDA_JURISDICTION_KEYWORDS: dict[str, list[str]] = {
    "india": ["nda india", "nda - india", "india nda"],
    "us": ["nda us", "nda - us", "us nda"],
    "singapore": ["nda singapore", "nda - singapore", "singapore nda"],
}


async def _find_nda_template_row(
    jurisdiction: Jurisdiction,
    user_id: Optional[str] = None,
) -> Optional[IndexedDriveFile]:
    """Locate the indexed NDA template row, preferring user's own then admin.

    Scope rules: prefer the user's own synced file, fall back to admin/shared,
    never leak another user's private file. Tries general "NDA"-flavoured
    keywords first; falls back to jurisdiction-specific phrasing only if the
    general search misses (covers users who file templates under e.g.
    "NDA - India.docx").
    """
    from sqlalchemy import or_

    keywords: list[str] = list(_NDA_GENERAL_KEYWORDS) + list(
        _NDA_JURISDICTION_KEYWORDS.get(jurisdiction, [])
    )

    async with async_session() as session:
        for keyword in keywords:
            stmt = sm_select(IndexedDriveFile).where(
                IndexedDriveFile.name.ilike(f"%{keyword}%"),
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        IndexedDriveFile.owner_user_id == user_id,
                        IndexedDriveFile.is_admin == True,  # noqa: E712
                    )
                )
            else:
                stmt = stmt.where(IndexedDriveFile.is_admin == True)  # noqa: E712

            if user_id is not None:
                from sqlalchemy import case
                user_priority = case(
                    (IndexedDriveFile.owner_user_id == user_id, 0),
                    else_=1,
                )
                stmt = stmt.order_by(
                    user_priority,
                    IndexedDriveFile.last_indexed_at.desc(),
                )
            else:
                stmt = stmt.order_by(IndexedDriveFile.last_indexed_at.desc())

            row = (await session.execute(stmt.limit(1))).scalar_one_or_none()
            if row:
                logger.info(
                    "NDA template found via keyword=%r jurisdiction=%s: "
                    "name=%s file_id=%s mime=%s owner=%s is_admin=%s",
                    keyword, jurisdiction, row.name, row.drive_file_id,
                    row.mime_type, row.owner_user_id, row.is_admin,
                )
                return row
    logger.info(
        "No NDA template matched for jurisdiction=%s user_id=%s",
        jurisdiction, user_id,
    )
    return None


async def _fetch_template_bytes(
    jurisdiction: Jurisdiction,
    user_id: Optional[str] = None,
) -> Optional[bytes]:
    """Download the NDA template, auto-discovering via IndexedDriveFile first.

    Order of attempts:
      1. Look up the file in IndexedDriveFile by keyword (user-scoped then
         admin-scoped). If found, fetch via the file owner's OAuth token and
         export Google Docs to .docx as needed.
      2. Fall back to the env-var-configured file_id and the admin connection.

    Returns None on any failure so the caller can decide whether to error out
    or render the fallback draft.
    """
    from app.clients import google_drive

    # --- 1. Auto-discovery via IndexedDriveFile -----------------------------
    row = await _find_nda_template_row(jurisdiction, user_id=user_id)
    if row:
        async with async_session() as session:
            result = await session.execute(
                sm_select(UserEmailConnection).where(
                    UserEmailConnection.user_id == row.owner_user_id,
                    UserEmailConnection.is_active == True,  # noqa: E712
                )
            )
            connection = result.scalar_one_or_none()

        if connection:
            try:
                if row.mime_type == "application/vnd.google-apps.document":
                    # Google Doc — export to .docx via Drive's export endpoint.
                    import httpx
                    from app.clients.google_drive import _ensure_token, DRIVE_API_BASE
                    access_token, _updated = await _ensure_token(
                        connection.token_data,
                        settings.gmail_client_id,
                        settings.gmail_client_secret,
                    )
                    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    async with httpx.AsyncClient(timeout=60) as http:
                        resp = await http.get(
                            f"{DRIVE_API_BASE}/files/{row.drive_file_id}/export",
                            headers={"Authorization": f"Bearer {access_token}"},
                            params={"mimeType": docx_mime},
                        )
                    if resp.status_code == 200:
                        data = resp.content
                    else:
                        logger.warning(
                            "NDA template Drive export failed (status %s): %s",
                            resp.status_code, resp.text[:300],
                        )
                        data = None
                else:
                    data, _mime, _updated = await google_drive.download_file_bytes(
                        file_id=row.drive_file_id,
                        mime_type=row.mime_type,
                        token_data=connection.token_data,
                        client_id=settings.gmail_client_id,
                        client_secret=settings.gmail_client_secret,
                    )

                if data and data.startswith(b"PK"):
                    logger.info(
                        "NDA template fetched via auto-discovery: %d bytes",
                        len(data),
                    )
                    return data
                if data:
                    logger.warning(
                        "NDA template '%s' didn't return a valid .docx "
                        "(first bytes: %r) — falling back to env-var template.",
                        row.name, data[:8],
                    )
            except Exception as exc:
                logger.exception(
                    "Auto-discovery fetch failed for NDA template '%s': %s",
                    row.name, exc,
                )
        else:
            logger.warning(
                "No active Drive connection for owner of NDA template '%s' — "
                "falling back to env-var template.",
                row.name,
            )

    # --- 2. Env-var fallback ------------------------------------------------
    file_id = _template_drive_id(jurisdiction)
    if not file_id:
        logger.info(
            "No Drive template configured for jurisdiction=%s and "
            "auto-discovery missed — caller will use synthetic fallback.",
            jurisdiction,
        )
        return None

    async with async_session() as session:
        result = await session.execute(
            sm_select(UserEmailConnection).where(
                UserEmailConnection.is_admin_folder == True,  # noqa: E712
                UserEmailConnection.is_active == True,  # noqa: E712
            )
        )
        connection = result.scalar_one_or_none()
        if not connection:
            result = await session.execute(
                sm_select(UserEmailConnection).where(
                    UserEmailConnection.is_active == True,  # noqa: E712
                )
            )
            connection = result.scalar_one_or_none()
        if not connection:
            logger.warning("No Drive connection available to fetch NDA template")
            return None

        try:
            data, _mime, _updated = await google_drive.download_file_bytes(
                file_id=file_id,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                token_data=connection.token_data,
                client_id=settings.gmail_client_id,
                client_secret=settings.gmail_client_secret,
            )
            return data
        except Exception as exc:
            logger.exception("Failed to download NDA template %s: %s", file_id, exc)
            return None


# ── Inspection ───────────────────────────────────────────────────────────────


async def inspect_template(
    jurisdiction: Jurisdiction,
    user_id: Optional[str] = None,
) -> dict:
    """Report template availability + a section summary for the given jurisdiction.

    The old version returned indexed blanks for the agent to fill; the new one
    simply tells the agent "template exists, it has N rewritable sections,
    go collect the party details." The rewrite engine handles the rest.
    """
    template_bytes = await _fetch_template_bytes(jurisdiction, user_id=user_id)
    if not template_bytes:
        return {
            "found": False,
            "jurisdiction": jurisdiction,
            "error": (
                "No NDA template found. Upload a file named 'Template of "
                "Beacon NDA.docx' (or any name containing 'NDA') to your "
                "indexed Drive folder and re-sync."
            ),
            "section_count": 0,
            "sections": [],
        }

    structure = extract_docx_structure(template_bytes)
    content_blocks = [b for b in structure if not b["is_structural"] and b.get("text", "").strip()]
    preview = [
        {"index": b["block_index"], "text": (b.get("text") or "")[:120]}
        for b in content_blocks[:10]
    ]
    return {
        "found": True,
        "jurisdiction": jurisdiction,
        "section_count": len(content_blocks),
        "sections": preview,
    }


# ── Fallback draft (only if template unavailable) ────────────────────────────


def _render_fallback(data: NDAInput, path) -> None:
    """Produce a fully-filled NDA .docx from the user-provided values.

    Used only when the Drive template can't be fetched. A visible banner makes
    clear that this is *not* the canonical template and needs counsel review.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    disclosing = data.disclosing_party or "RARE BITS TECHNOLOGY PRIVATE LIMITED"
    receiving = data.receiving_party or "____"
    effective = data.effective_date or "____"
    city = data.governing_city or "____"
    term = str(data.term_years) if data.term_years is not None else "____"
    purpose = data.purpose or "the discussion of a potential business relationship"
    mutual = bool(data.mutual) if data.mutual is not None else True

    doc = Document()

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = banner.add_run(
        "DRAFT — No NDA template found in Drive. Upload a file named "
        "'Template of Beacon NDA.docx' (or any name containing 'NDA') to "
        "your indexed Drive folder and re-sync. Have counsel review before "
        "execution."
    )
    br.italic = True
    br.font.size = Pt(9)
    br.font.color.rgb = RGBColor(0xB2, 0x00, 0x00)

    title = doc.add_heading(
        "MUTUAL NON-DISCLOSURE AGREEMENT" if mutual else "NON-DISCLOSURE AGREEMENT",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"This Non-Disclosure Agreement (the \"Agreement\") is entered into on "
        f"{effective} by and between {disclosing}, a company incorporated under "
        f"the laws of the {_JURISDICTION_LABEL[data.jurisdiction]} (the "
        f"\"Disclosing Party\"), and {receiving} (the \"Receiving Party\"). "
        f"The Disclosing Party and the Receiving Party are each a \"Party\" "
        f"and collectively the \"Parties\"."
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        f"The Parties wish to explore {purpose} (the \"Purpose\") and, in "
        f"connection therewith, may disclose to each other certain "
        f"confidential and proprietary information."
    )

    doc.add_heading("2. Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means any non-public information "
        "disclosed by one Party to the other, whether orally, in writing, or "
        "by inspection of tangible objects, that is designated as confidential "
        "or that reasonably should be understood to be confidential given the "
        "nature of the information and the circumstances of disclosure."
    )

    doc.add_heading("3. Obligations", level=1)
    doc.add_paragraph(
        "The Receiving Party shall (i) hold all Confidential Information in "
        "strict confidence; (ii) use such information solely for the Purpose; "
        "and (iii) not disclose it to any third party without the prior written "
        "consent of the Disclosing Party."
    )

    doc.add_heading("4. Term", level=1)
    doc.add_paragraph(
        f"This Agreement shall remain in effect for a period of {term} year(s) "
        f"from the Effective Date, and the obligations of confidentiality "
        f"shall survive termination for a further period of two (2) years."
    )

    doc.add_heading("5. Governing Law", level=1)
    doc.add_paragraph(
        f"This Agreement shall be governed by the laws of the "
        f"{_JURISDICTION_LABEL[data.jurisdiction]}, and the courts at {city} "
        f"shall have exclusive jurisdiction."
    )

    doc.add_heading("6. Signatures", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = f"{disclosing}"
    table.rows[0].cells[1].text = f"{receiving}"
    table.rows[1].cells[0].text = "By: ______________________"
    table.rows[1].cells[1].text = "By: ______________________"
    table.rows[2].cells[0].text = "Name: ____________________"
    table.rows[2].cells[1].text = "Name: ____________________"
    table.rows[3].cells[0].text = "Title: ___________________"
    table.rows[3].cells[1].text = "Title: ___________________"

    doc.save(str(path))


# ── Public entry ─────────────────────────────────────────────────────────────


async def generate(data: NDAInput, user_id: Optional[str] = None) -> GeneratedDocument:
    """Generate an NDA by rewriting the Drive template with the user's inputs."""
    if data.jurisdiction not in _JURISDICTION_LABEL:
        raise ValueError(
            f"Unsupported NDA jurisdiction: {data.jurisdiction}. "
            "Supported: india, us, singapore."
        )

    slug_left = data.disclosing_party or "Beacon"
    slug_right = data.receiving_party or "counterparty"
    path, url = build_output_path(
        f"NDA_{data.jurisdiction.upper()}",
        f"{slug_left}_x_{slug_right}",
    )

    template_bytes = await _fetch_template_bytes(data.jurisdiction, user_id=user_id)
    source = "template"

    if not template_bytes:
        logger.warning(
            "NDA template for %s unreachable — rendering fallback draft.",
            data.jurisdiction,
        )
        _render_fallback(data, path)
        source = "fallback"
    else:
        structure = extract_docx_structure(template_bytes)

        user_inputs = {
            "jurisdiction": _JURISDICTION_LABEL[data.jurisdiction],
            "disclosing_party": data.disclosing_party or "",
            "receiving_party": data.receiving_party or "",
            "effective_date": data.effective_date or "",
            "governing_city": data.governing_city or "",
            "term_years": data.term_years if data.term_years is not None else "",
            "purpose": data.purpose or "",
            "mutual": bool(data.mutual) if data.mutual is not None else None,
            "extra_clauses": data.extra_clauses or [],
            # Pass legacy ``fills`` through verbatim — Claude can use any
            # extra context the agent collected (registered office, PAN, etc.).
            "additional_details": data.fills or {},
        }

        client = get_anthropic_client()
        rewritten = await rewrite_with_claude(
            structure=structure,
            user_inputs=user_inputs,
            doc_type="nda",
            client=client,
            model=settings.CLAUDE_MODEL_STANDARD,
        )

        out_bytes = rewrite_docx_content(template_bytes, rewritten)
        path.write_bytes(out_bytes)

    result = GeneratedDocument(
        filename=path.name,
        path=str(path),
        url=url,
        kind=f"nda_{data.jurisdiction}",
        summary=(
            f"NDA drafted from {source} — {slug_left} ↔ {slug_right}, "
            f"{data.jurisdiction.upper()} jurisdiction."
        ),
        created_at=datetime.utcnow(),
    )

    await _try_upload_to_drive(result, path, f"{slug_left} x {slug_right}", user_id=user_id)
    return result


async def _try_upload_to_drive(
    doc: GeneratedDocument,
    path,
    label: str,
    *,
    user_id: Optional[str],
) -> None:
    """Upload the filled NDA .docx to the user's Drive folder as a Google Doc."""
    if doc.drive_url:
        return
    if user_id:
        from app.services.zippy_docs.base import (
            cache_upload,
            get_cached_upload,
        )
        cached = get_cached_upload(str(user_id), label, doc.kind)
        if cached:
            doc.drive_url = cached
            logger.info("Reusing cached NDA Drive upload: %s", cached)
            return
    try:
        async with async_session() as session:
            from sqlalchemy import or_
            stmt = sm_select(UserEmailConnection).where(
                UserEmailConnection.is_active == True,  # noqa: E712
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        UserEmailConnection.user_id == user_id,
                        UserEmailConnection.is_admin_folder == True,  # noqa: E712
                    )
                )
            from sqlalchemy import case as sa_case
            if user_id is not None:
                priority = sa_case(
                    (UserEmailConnection.user_id == user_id, 0), else_=1
                )
                stmt = stmt.order_by(priority)
            result = await session.execute(stmt.limit(1))
            connection = result.scalar_one_or_none()

        if not connection:
            logger.info("No active Drive connection — skipping Google Docs upload for NDA")
            return

        docx_bytes = path.read_bytes()
        gdoc_name = f"NDA — {label} — {doc.created_at.strftime('%d %b %Y')}"
        folder_id = connection.selected_drive_folder_id or None

        file_id, web_view_link = await upload_as_google_doc(
            filename=gdoc_name,
            docx_bytes=docx_bytes,
            token_data=connection.token_data,
            client_id=settings.gmail_client_id,
            client_secret=settings.gmail_client_secret,
            parent_folder_id=folder_id,
        )
        doc.drive_file_id = file_id
        doc.drive_url = web_view_link
        logger.info("NDA uploaded to Google Docs: %s", web_view_link)
        if user_id and web_view_link:
            from app.services.zippy_docs.base import cache_upload
            cache_upload(str(user_id), label, doc.kind, web_view_link)
    except PermissionError as exc:
        logger.info("drive.file scope not yet granted — skipping NDA upload: %s", exc)
    except Exception as exc:
        logger.warning("Google Docs upload failed for NDA (non-fatal): %s", exc)
