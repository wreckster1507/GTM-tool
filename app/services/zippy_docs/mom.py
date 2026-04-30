"""Minutes of Meeting generator — fills Beacon's actual Drive MOM template.

New flow (no {{TOKEN}} placeholders required):
  1. ``inspect_mom_template()`` downloads the Drive template and reports how
     many content sections it contains — the agent uses this to confirm the
     template is available before asking for a transcript.
  2. ``generate(data)`` downloads the template, extracts every paragraph as a
     structured block (flagging headings/labels as "structural"), asks Claude
     to rewrite only the non-structural content blocks from the transcript,
     and patches the rewrites back in while preserving run-level formatting.

If the Drive template can't be fetched, ``generate`` still produces a usable
MOM via ``_render_fallback_docx`` so the tool never hard-fails.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from sqlmodel import select as sm_select

from app.clients.anthropic_client import get_anthropic_client
from app.clients.google_drive import upload_as_google_doc
from app.config import settings
from app.database import AsyncSessionLocal as async_session
from app.models.user_email_connection import UserEmailConnection
from app.models.zippy import IndexedDriveFile
from app.services.zippy_docs.base import (
    GeneratedDocument,
    build_output_path,
    human_today,
)
from app.services.zippy_docs.claude_rewriter import rewrite_with_claude
from app.services.zippy_docs.doc_rewriter import (
    extract_docx_structure,
    rewrite_docx_content,
)

logger = logging.getLogger(__name__)


@dataclass
class MOMInput:
    client_name: str
    meeting_date: Optional[str] = None
    attendees: Optional[list[str]] = None
    transcript: Optional[str] = None
    context_notes: Optional[str] = None
    # "long" = detailed (all sections), "short" = key highlights only
    format_type: str = "long"
    # Collateral list the agent determined should be included — each item is
    # a string like "Deck : Beacon – Implementation Automation | <url>"
    collateral: Optional[list[str]] = None


# ── Drive template fetch — auto-discovers from indexed_drive_files ────────────
#
# Unchanged from the pre-refactor version. The heavy lifting (ownership
# filtering, Google-Doc-to-.docx export, zip magic validation, per-file cache)
# still lives here because the rewrite engine is doc-type agnostic — it
# doesn't know how to locate a template, only how to transform one.

_MOM_KEYWORDS = ["mom template", "mom", "minutes of meeting"]

_TEMPLATE_CACHE: dict[str, tuple[str, bytes]] = {}


async def _find_mom_template_row(user_id: Optional[str] = None) -> Optional[IndexedDriveFile]:
    """Return the IndexedDriveFile row for the MOM template, or None.

    Scope rules: prefer the user's own synced file, fall back to admin/shared,
    never leak another user's private file.
    """
    from sqlalchemy import or_

    async with async_session() as session:
        for keyword in _MOM_KEYWORDS:
            stmt = sm_select(IndexedDriveFile).where(
                IndexedDriveFile.name.ilike(f"%{keyword}%"),
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        IndexedDriveFile.owner_user_id == user_id,
                        IndexedDriveFile.is_admin == True,  # noqa: E712
                    )
                )
            else:
                stmt = stmt.where(IndexedDriveFile.is_admin == True)  # noqa: E712

            if user_id is not None:
                from sqlalchemy import case
                user_priority = case(
                    (IndexedDriveFile.owner_user_id == user_id, 0),
                    else_=1,
                )
                stmt = stmt.order_by(
                    user_priority,
                    IndexedDriveFile.last_indexed_at.desc(),
                )
            else:
                stmt = stmt.order_by(IndexedDriveFile.last_indexed_at.desc())

            row = (await session.execute(stmt.limit(1))).scalar_one_or_none()
            if row:
                logger.info(
                    "MOM template found: name=%s file_id=%s mime=%s owner=%s is_admin=%s",
                    row.name, row.drive_file_id, row.mime_type,
                    row.owner_user_id, row.is_admin,
                )
                return row
    logger.info("No MOM template matched for user_id=%s", user_id)
    return None


async def _fetch_template_bytes_with_error(
    user_id: Optional[str] = None,
) -> tuple[Optional[bytes], Optional[str]]:
    """Fetch the MOM template bytes for a specific user.

    Uses the OAuth connection that matches the file's owner so personal files
    go through the user's token and admin-shared files go through the admin's.
    """
    row = await _find_mom_template_row(user_id=user_id)
    if not row:
        return None, (
            "No MOM template is indexed for this user. Upload a file named "
            "'MOM Template.docx' to your synced Drive folder and click Sync."
        )

    file_id = row.drive_file_id
    actual_mime = row.mime_type
    cache_key = (
        row.drive_modified_at.isoformat() if row.drive_modified_at else "no-modified"
    )

    cached = _TEMPLATE_CACHE.get(file_id)
    if cached and cached[0] == cache_key:
        logger.info("MOM template cache HIT for %s (%d bytes)", file_id, len(cached[1]))
        return cached[1], None

    from app.clients import google_drive

    async with async_session() as session:
        result = await session.execute(
            sm_select(UserEmailConnection).where(
                UserEmailConnection.user_id == row.owner_user_id,
                UserEmailConnection.is_active == True,  # noqa: E712
            )
        )
        connection = result.scalar_one_or_none()

        if not connection:
            logger.warning(
                "No active Drive connection for owner_user_id=%s — cannot fetch MOM template",
                row.owner_user_id,
            )
            return None, (
                f"Drive connection for the owner of '{row.name}' is inactive. "
                "Reconnect Google Drive in Settings and re-sync."
            )

        try:
            if actual_mime == "application/vnd.google-apps.document":
                import httpx
                from app.clients.google_drive import _ensure_token, DRIVE_API_BASE
                access_token, _updated_token = await _ensure_token(
                    connection.token_data,
                    settings.gmail_client_id,
                    settings.gmail_client_secret,
                )
                docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                async with httpx.AsyncClient(timeout=60) as http:
                    resp = await http.get(
                        f"{DRIVE_API_BASE}/files/{file_id}/export",
                        headers={"Authorization": f"Bearer {access_token}"},
                        params={"mimeType": docx_mime},
                    )
                if resp.status_code != 200:
                    return None, (
                        f"Drive export to .docx failed for '{row.name}' "
                        f"(status {resp.status_code}): {resp.text[:300]}"
                    )
                data = resp.content
            else:
                data, _mime, _updated = await google_drive.download_file_bytes(
                    file_id=file_id,
                    mime_type=actual_mime,
                    token_data=connection.token_data,
                    client_id=settings.gmail_client_id,
                    client_secret=settings.gmail_client_secret,
                )
            if not data:
                return None, (
                    f"Drive returned empty bytes for {row.name} (file_id={file_id}, "
                    f"mime={actual_mime}). File may have been deleted or permissions revoked."
                )
            if not data.startswith(b"PK"):
                return None, (
                    f"'{row.name}' didn't come back as a valid .docx "
                    f"(first bytes: {data[:8]!r}). Re-upload the file as a "
                    "real .docx (don't let Drive convert it to a Google Doc)."
                )
            logger.info("MOM template downloaded fresh: %d bytes", len(data))
            _TEMPLATE_CACHE[file_id] = (cache_key, data)
            return data, None
        except Exception as exc:
            logger.exception("Failed to download MOM template file_id=%s: %s", file_id, exc)
            return None, f"Drive download failed for {row.name}: {type(exc).__name__}: {exc}"


# ── Inspection ───────────────────────────────────────────────────────────────


async def inspect_mom_template(user_id: Optional[str] = None) -> dict:
    """Report what the MOM template looks like so the agent can decide next steps.

    Replaces the old "list every {{TOKEN}}" inspection with a sections summary:
    how many content (non-structural) blocks the template has, plus a small
    preview. The agent doesn't need to map tokens anymore — it just needs to
    know the template exists and has content Claude can rewrite.
    """
    template_bytes, err = await _fetch_template_bytes_with_error(user_id=user_id)
    if not template_bytes:
        return {
            "found": False,
            "error": err or "Unknown error fetching MOM template.",
            "section_count": 0,
            "sections": [],
        }

    structure = extract_docx_structure(template_bytes)
    content_blocks = [b for b in structure if not b["is_structural"] and b.get("text", "").strip()]
    # Small preview: first 10 content blocks, text truncated. Keeps the tool
    # result compact so it stays cheap to send back through the model loop.
    preview = [
        {"index": b["block_index"], "text": (b.get("text") or "")[:120]}
        for b in content_blocks[:10]
    ]

    row = await _find_mom_template_row(user_id=user_id)
    return {
        "found": True,
        "template_name": row.name if row else None,
        "template_drive_file_id": row.drive_file_id if row else None,
        "section_count": len(content_blocks),
        "sections": preview,
    }


# ── Fallback renderer (no template available) ────────────────────────────────


def _render_fallback_docx(data: MOMInput, path) -> None:
    """Build a standard MOM .docx when the Drive template can't be fetched.

    This is intentionally plain — the whole point of the template flow is
    brand-consistent output, so we only land here when Drive is unreachable or
    the user hasn't indexed a template yet.
    """
    doc = Document()
    title = doc.add_heading(f"Minutes of Meeting — {data.client_name}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)

    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run(data.meeting_date or human_today())
    meta.add_run("\n")
    meta.add_run("Attendees: ").bold = True
    meta.add_run(", ".join(data.attendees or []) or "Not specified")

    def heading(text: str) -> None:
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x17, 0x4A, 0x8B)

    heading("Overview")
    body = (data.transcript or data.context_notes or "").strip()
    if body:
        for para in body.split("\n\n"):
            doc.add_paragraph(para.strip())
    else:
        p = doc.add_paragraph("No transcript or notes were provided.")
        p.runs[0].italic = True

    footer = doc.add_paragraph()
    fr = footer.add_run(
        f"\nGenerated by Zippy on {datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}"
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    doc.save(str(path))


# ── Public entry points ───────────────────────────────────────────────────────


class MOMTemplateUnavailable(RuntimeError):
    """Kept for backwards compatibility with callers that imported it.

    The new ``generate()`` no longer raises this — it falls back to the
    plain-template renderer so the agent can always hand the user *something*.
    """


async def generate(data: MOMInput, user_id: Optional[str] = None) -> GeneratedDocument:
    """Generate a MOM by rewriting the Drive template with transcript content.

    Steps:
      1. Fetch template bytes. If unavailable → render fallback docx.
      2. Extract doc structure (every paragraph, tagged structural vs content).
      3. Ask Claude to rewrite the content blocks from ``data``.
      4. Patch rewrites back into the original bytes, preserving formatting.
      5. Upload to Drive (best-effort) and return the GeneratedDocument.
    """
    path, url = build_output_path("MOM", data.client_name)
    template_bytes, err = await _fetch_template_bytes_with_error(user_id=user_id)
    source = "template"

    if not template_bytes:
        logger.warning("MOM template unavailable (%s) — rendering fallback docx", err)
        _render_fallback_docx(data, path)
        source = "fallback"
    else:
        # Extract → rewrite → patch.
        structure = extract_docx_structure(template_bytes)
        # Visibility into the structural-vs-content split — if "empty slots"
        # is 0 on a template that visibly has empty content rows, the old
        # `_looks_structural` is still active and the backend needs reloading.
        empty_slots = sum(
            1 for b in structure
            if not b["is_structural"] and not (b.get("text") or "").strip()
        )
        non_empty_content = sum(
            1 for b in structure
            if not b["is_structural"] and (b.get("text") or "").strip()
        )
        logger.info(
            "MOM template structure: %d total blocks, %d structural, "
            "%d non-empty content, %d empty slots",
            len(structure),
            sum(1 for b in structure if b["is_structural"]),
            non_empty_content,
            empty_slots,
        )

        user_inputs = {
            "client_name": data.client_name,
            "meeting_date": data.meeting_date or human_today(),
            "attendees": data.attendees or [],
            "transcript": (data.transcript or data.context_notes or "")[:20000],
            "format_type": data.format_type or "long",
            "collateral": data.collateral or [],
        }

        client = get_anthropic_client()
        rewritten = await rewrite_with_claude(
            structure=structure,
            user_inputs=user_inputs,
            doc_type="mom",
            client=client,
            model=settings.CLAUDE_MODEL_STANDARD,
        )

        out_bytes = rewrite_docx_content(template_bytes, rewritten)
        path.write_bytes(out_bytes)

    result = GeneratedDocument(
        filename=path.name,
        path=str(path),
        url=url,
        kind="mom",
        summary=(
            f"MOM drafted from {source} for {data.client_name}."
        ),
        created_at=datetime.utcnow(),
    )

    await _try_upload_to_drive(result, path, data.client_name, user_id=user_id)
    return result


async def _try_upload_to_drive(
    doc: GeneratedDocument,
    path,
    client_name: str,
    *,
    user_id: Optional[str],
) -> None:
    """Upload the filled .docx to the user's Drive folder as a Google Doc.

    Unchanged from pre-refactor — best-effort, silently skipped on failure.
    """
    try:
        async with async_session() as session:
            from sqlalchemy import or_
            stmt = sm_select(UserEmailConnection).where(
                UserEmailConnection.is_active == True,  # noqa: E712
            )
            if user_id is not None:
                stmt = stmt.where(
                    or_(
                        UserEmailConnection.user_id == user_id,
                        UserEmailConnection.is_admin_folder == True,  # noqa: E712
                    )
                )
            from sqlalchemy import case as sa_case
            if user_id is not None:
                priority = sa_case(
                    (UserEmailConnection.user_id == user_id, 0), else_=1
                )
                stmt = stmt.order_by(priority)
            result = await session.execute(stmt.limit(1))
            connection = result.scalar_one_or_none()

        if not connection:
            logger.info("No active Drive connection found — skipping Google Docs upload")
            return

        docx_bytes = path.read_bytes()
        gdoc_name = f"MOM — {client_name} — {doc.created_at.strftime('%d %b %Y')}"
        folder_id = connection.selected_drive_folder_id or None

        file_id, web_view_link = await upload_as_google_doc(
            filename=gdoc_name,
            docx_bytes=docx_bytes,
            token_data=connection.token_data,
            client_id=settings.gmail_client_id,
            client_secret=settings.gmail_client_secret,
            parent_folder_id=folder_id,
        )
        doc.drive_file_id = file_id
        doc.drive_url = web_view_link
        logger.info("MOM uploaded to Google Docs: %s", web_view_link)
    except PermissionError as exc:
        logger.info("drive.file scope not yet granted — skipping upload: %s", exc)
    except Exception as exc:
        logger.warning("Google Docs upload failed (non-fatal): %s", exc)
